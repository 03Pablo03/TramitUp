"""
Servicio de plantillas inteligentes con generación condicional.
Crea documentos adaptativos basados en contexto, entidades extraídas y personalización.
"""

import json
import re
from datetime import datetime, date
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.ai.llm_client import get_llm
from app.services.legal_entity_extractor import legal_entity_extractor
from app.services.personalization_service import personalization_service


@dataclass
class TemplateCondition:
    """Condición para mostrar/ocultar secciones de plantilla."""
    field: str  # Campo a evaluar
    operator: str  # 'equals', 'contains', 'exists', 'not_exists', 'greater_than', 'less_than'
    value: Any  # Valor a comparar
    section_id: str  # ID de la sección afectada


@dataclass
class TemplateSection:
    """Sección modular de una plantilla."""
    id: str
    title: str
    content: str
    conditions: List[TemplateCondition]
    required: bool = False
    order: int = 0


@dataclass
class SmartTemplate:
    """Plantilla inteligente con secciones modulares."""
    id: str
    name: str
    description: str
    document_type: str
    sections: List[TemplateSection]
    global_fields: Dict[str, Any]
    ai_enhancement: bool = True


class IntelligentTemplateService:
    """Servicio de plantillas inteligentes con generación condicional."""
    
    def __init__(self):
        self.templates = self._load_smart_templates()
        self.field_extractors = self._setup_field_extractors()
    
    def generate_adaptive_document(
        self,
        template_id: str,
        user_context: Dict[str, Any],
        extracted_entities: Dict[str, Any],
        conversation_context: str,
        user_id: str
    ) -> Dict[str, Any]:
        """
        Genera un documento adaptativo usando plantillas inteligentes.
        
        Args:
            template_id: ID de la plantilla a usar
            user_context: Contexto del usuario (personalización)
            extracted_entities: Entidades extraídas de documentos
            conversation_context: Contexto de la conversación
            user_id: ID del usuario
            
        Returns:
            Diccionario con el documento generado y metadatos
        """
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Plantilla no encontrada: {template_id}")
        
        # 1. Extraer campos dinámicamente
        dynamic_fields = self._extract_dynamic_fields(
            template, user_context, extracted_entities, conversation_context
        )
        
        # 2. Evaluar condiciones y filtrar secciones
        active_sections = self._evaluate_template_conditions(
            template, dynamic_fields, user_context
        )
        
        # 3. Enriquecer contenido con IA si está habilitado
        if template.ai_enhancement:
            active_sections = self._enhance_sections_with_ai(
                active_sections, dynamic_fields, conversation_context, user_context
            )
        
        # 4. Generar documento Word
        doc_path = self._generate_word_document(
            template, active_sections, dynamic_fields, user_id
        )
        
        return {
            "document_path": doc_path,
            "template_used": template.name,
            "sections_included": [s.title for s in active_sections],
            "fields_populated": len([k for k, v in dynamic_fields.items() if v]),
            "ai_enhanced": template.ai_enhancement,
            "generation_metadata": {
                "total_sections_available": len(template.sections),
                "sections_activated": len(active_sections),
                "conditions_evaluated": sum(len(s.conditions) for s in template.sections),
                "user_personalization_applied": bool(user_context),
                "entities_used": len(extracted_entities.get("entities", [])),
                "generation_timestamp": datetime.now().isoformat()
            }
        }
    
    def _load_smart_templates(self) -> Dict[str, SmartTemplate]:
        """Carga plantillas inteligentes predefinidas."""
        templates = {}
        
        # Plantilla: Reclamación de Factura Inteligente
        templates["reclamacion_factura_smart"] = SmartTemplate(
            id="reclamacion_factura_smart",
            name="Reclamación de Factura Inteligente",
            description="Reclamación adaptativa que se ajusta según el tipo de factura y entidades detectadas",
            document_type="reclamacion",
            sections=[
                TemplateSection(
                    id="header",
                    title="Encabezado",
                    content="""
{nombre_completo}
{direccion_completa}
{telefono}
{email}

{ciudad}, {fecha_actual}

{empresa_destinataria}
{direccion_empresa}
Att: Departamento de Atención al Cliente
                    """,
                    conditions=[],
                    required=True,
                    order=1
                ),
                TemplateSection(
                    id="asunto",
                    title="Asunto",
                    content="ASUNTO: Reclamación por {tipo_factura} - Factura nº {numero_factura}",
                    conditions=[
                        TemplateCondition("numero_factura", "exists", None, "asunto")
                    ],
                    required=True,
                    order=2
                ),
                TemplateSection(
                    id="introduccion_persona_fisica",
                    title="Introducción - Persona Física",
                    content="""
Estimados señores,

Me dirijo a ustedes como cliente particular con DNI {dni}, para presentar una reclamación formal respecto a la factura mencionada en el asunto.
                    """,
                    conditions=[
                        TemplateCondition("dni", "exists", None, "introduccion_persona_fisica"),
                        TemplateCondition("tipo_cliente", "equals", "particular", "introduccion_persona_fisica")
                    ],
                    required=False,
                    order=3
                ),
                TemplateSection(
                    id="introduccion_empresa",
                    title="Introducción - Empresa",
                    content="""
Estimados señores,

Me dirijo a ustedes en representación de {nombre_empresa}, con CIF {cif}, para presentar una reclamación formal respecto a la factura mencionada en el asunto.
                    """,
                    conditions=[
                        TemplateCondition("cif", "exists", None, "introduccion_empresa"),
                        TemplateCondition("tipo_cliente", "equals", "empresa", "introduccion_empresa")
                    ],
                    required=False,
                    order=3
                ),
                TemplateSection(
                    id="detalles_factura",
                    title="Detalles de la Factura",
                    content="""
DATOS DE LA FACTURA:
- Número de factura: {numero_factura}
- Fecha de emisión: {fecha_factura}
- Importe total: {importe_total}
- Período facturado: {periodo_facturado}
- Número de contrato/cliente: {numero_contrato}
                    """,
                    conditions=[],
                    required=True,
                    order=4
                ),
                TemplateSection(
                    id="motivo_importe_excesivo",
                    title="Motivo - Importe Excesivo",
                    content="""
MOTIVO DE LA RECLAMACIÓN:

Considero que el importe facturado de {importe_total} es excesivo e incorrecto por los siguientes motivos:

1. El consumo facturado ({consumo_facturado}) no se corresponde con mi consumo habitual
2. Se han aplicado tarifas incorrectas o conceptos no contratados
3. {motivo_personalizado}

Solicito la revisión detallada de la facturación y la emisión de una factura rectificativa.
                    """,
                    conditions=[
                        TemplateCondition("tipo_reclamacion", "equals", "importe_excesivo", "motivo_importe_excesivo")
                    ],
                    required=False,
                    order=5
                ),
                TemplateSection(
                    id="motivo_servicios_no_prestados",
                    title="Motivo - Servicios No Prestados",
                    content="""
MOTIVO DE LA RECLAMACIÓN:

Rechazo el pago de la presente factura ya que los servicios facturados no han sido prestados o han sido prestados de forma deficiente:

1. Período sin suministro: del {fecha_inicio_corte} al {fecha_fin_corte}
2. Incidencias reportadas: {incidencias_reportadas}
3. {motivo_personalizado}

Solicito la anulación de los importes correspondientes al período sin servicio.
                    """,
                    conditions=[
                        TemplateCondition("tipo_reclamacion", "equals", "servicios_no_prestados", "motivo_servicios_no_prestados")
                    ],
                    required=False,
                    order=5
                ),
                TemplateSection(
                    id="legislacion_consumo",
                    title="Marco Legal - Consumo",
                    content="""
FUNDAMENTOS LEGALES:

Esta reclamación se basa en la normativa de protección al consumidor:
- Real Decreto Legislativo 1/2007, Ley General para la Defensa de los Consumidores y Usuarios
- {normativa_sectorial_especifica}
- Derecho a recibir servicios de calidad y facturación correcta
                    """,
                    conditions=[
                        TemplateCondition("tipo_cliente", "equals", "particular", "legislacion_consumo")
                    ],
                    required=False,
                    order=6
                ),
                TemplateSection(
                    id="legislacion_mercantil",
                    title="Marco Legal - Mercantil",
                    content="""
FUNDAMENTOS LEGALES:

Esta reclamación se fundamenta en:
- Código Civil y Código de Comercio
- Condiciones contractuales acordadas
- {normativa_sectorial_especifica}
- Principio de buena fe contractual
                    """,
                    conditions=[
                        TemplateCondition("tipo_cliente", "equals", "empresa", "legislacion_mercantil")
                    ],
                    required=False,
                    order=6
                ),
                TemplateSection(
                    id="solicitudes",
                    title="Solicitudes",
                    content="""
Por todo lo expuesto, SOLICITO:

1. La revisión inmediata de la facturación objeto de esta reclamación
2. La emisión de una factura rectificativa con el importe correcto
3. La devolución del importe cobrado en exceso: {importe_reclamado}
4. Confirmación por escrito de las medidas adoptadas
5. {solicitud_adicional}

Plazo para respuesta: {plazo_respuesta} días hábiles conforme a la normativa aplicable.
                    """,
                    conditions=[],
                    required=True,
                    order=7
                ),
                TemplateSection(
                    id="advertencia_organismos",
                    title="Advertencia - Organismos Reguladores",
                    content="""
ADVERTENCIA:

En caso de no recibir respuesta satisfactoria en el plazo indicado, me reservo el derecho de:
- Presentar reclamación ante {organismo_regulador}
- Acudir a las Oficinas Municipales de Información al Consumidor (OMIC)
- Ejercitar las acciones legales que correspondan
                    """,
                    conditions=[
                        TemplateCondition("incluir_advertencia", "equals", True, "advertencia_organismos")
                    ],
                    required=False,
                    order=8
                ),
                TemplateSection(
                    id="cierre",
                    title="Cierre",
                    content="""
Sin otro particular, quedo a la espera de su pronta respuesta.

Atentamente,

{nombre_completo}
{dni_cif}

Adjunto: {documentos_adjuntos}
                    """,
                    conditions=[],
                    required=True,
                    order=9
                )
            ],
            global_fields={
                "fecha_actual": datetime.now().strftime("%d de %B de %Y"),
                "plazo_respuesta": "30"
            },
            ai_enhancement=True
        )
        
        # Plantilla: Recurso de Multa Inteligente
        templates["recurso_multa_smart"] = SmartTemplate(
            id="recurso_multa_smart",
            name="Recurso de Multa Inteligente",
            description="Recurso adaptativo según tipo de infracción y circunstancias",
            document_type="recurso",
            sections=[
                TemplateSection(
                    id="header",
                    title="Encabezado",
                    content="""
{nombre_completo}
{dni}
{direccion_completa}

{ciudad}, {fecha_actual}

{organismo_sancionador}
{direccion_organismo}
                    """,
                    conditions=[],
                    required=True,
                    order=1
                ),
                TemplateSection(
                    id="asunto",
                    title="Asunto",
                    content="ASUNTO: Recurso de Reposición contra Sanción nº {numero_expediente}",
                    conditions=[],
                    required=True,
                    order=2
                ),
                TemplateSection(
                    id="introduccion",
                    title="Introducción",
                    content="""
{nombre_completo}, con DNI {dni}, ante este organismo comparezco y, como mejor proceda en derecho, DIGO:

Que por medio del presente escrito vengo a interponer RECURSO DE REPOSICIÓN contra la resolución sancionadora de fecha {fecha_resolucion}, por la que se me impone una sanción de {importe_multa} euros.
                    """,
                    conditions=[],
                    required=True,
                    order=3
                ),
                TemplateSection(
                    id="alegacion_prescripcion",
                    title="Alegación - Prescripción",
                    content="""
PRIMERA.- PRESCRIPCIÓN DE LA INFRACCIÓN

Los hechos supuestamente constitutivos de infracción tuvieron lugar el {fecha_infraccion}, habiendo transcurrido más de {tiempo_transcurrido} desde entonces hasta la incoación del expediente el {fecha_incoacion}.

Conforme al artículo {articulo_prescripcion} de {normativa_aplicable}, el plazo de prescripción para este tipo de infracciones es de {plazo_prescripcion}, por lo que la infracción ha prescrito.
                    """,
                    conditions=[
                        TemplateCondition("alegacion_prescripcion", "equals", True, "alegacion_prescripcion")
                    ],
                    required=False,
                    order=4
                ),
                TemplateSection(
                    id="alegacion_defectos_procedimiento",
                    title="Alegación - Defectos de Procedimiento",
                    content="""
SEGUNDA.- DEFECTOS EN EL PROCEDIMIENTO

El procedimiento sancionador adolece de los siguientes defectos:

1. {defecto_notificacion}
2. {defecto_tramitacion}
3. Vulneración del derecho de defensa
4. {otros_defectos}

Estos defectos determinan la nulidad del procedimiento conforme a los artículos 47 y 48 de la Ley 39/2015.
                    """,
                    conditions=[
                        TemplateCondition("alegacion_defectos", "equals", True, "alegacion_defectos_procedimiento")
                    ],
                    required=False,
                    order=5
                ),
                TemplateSection(
                    id="alegacion_fondo",
                    title="Alegación de Fondo",
                    content="""
TERCERA.- ALEGACIONES DE FONDO

En cuanto al fondo del asunto, debo manifestar que:

1. Los hechos imputados no se ajustan a la realidad
2. {circunstancias_exculpatorias}
3. No concurren los elementos típicos de la infracción
4. {alegaciones_especificas}

Por tanto, no existe responsabilidad sancionadora alguna.
                    """,
                    conditions=[],
                    required=True,
                    order=6
                ),
                TemplateSection(
                    id="solicitudes",
                    title="Solicitudes",
                    content="""
Por todo lo expuesto, SOLICITO:

1. Se tenga por presentado este recurso de reposición
2. Se declare la {motivo_estimacion} del expediente
3. Se proceda al archivo de las actuaciones
4. Subsidiariamente, se reduzca la sanción aplicando los criterios de graduación

Todo ello con imposición de costas a la Administración.
                    """,
                    conditions=[],
                    required=True,
                    order=7
                )
            ],
            global_fields={
                "fecha_actual": datetime.now().strftime("%d de %B de %Y")
            },
            ai_enhancement=True
        )
        
        return templates
    
    def _setup_field_extractors(self) -> Dict[str, callable]:
        """Configura extractores de campos automáticos."""
        return {
            "dni": lambda entities: self._extract_entity_value(entities, "dni"),
            "cif": lambda entities: self._extract_entity_value(entities, "cif"),
            "numero_factura": lambda entities: self._extract_entity_value(entities, "factura"),
            "importe_total": lambda entities: self._extract_entity_value(entities, "euros"),
            "fecha_factura": lambda entities: self._extract_entity_value(entities, "fecha"),
            "telefono": lambda entities: self._extract_entity_value(entities, "telefono"),
            "email": lambda entities: self._extract_entity_value(entities, "email"),
            "direccion": lambda entities: self._extract_entity_value(entities, "direccion"),
        }
    
    def _extract_dynamic_fields(
        self,
        template: SmartTemplate,
        user_context: Dict[str, Any],
        extracted_entities: Dict[str, Any],
        conversation_context: str
    ) -> Dict[str, Any]:
        """Extrae campos dinámicamente de múltiples fuentes."""
        fields = template.global_fields.copy()
        
        # 1. Campos del perfil de usuario
        user_profile = user_context.get("profile", {})
        fields.update({
            "nombre_completo": user_profile.get("full_name", "[NOMBRE COMPLETO]"),
            "email": user_profile.get("email", "[EMAIL]"),
            "telefono": user_profile.get("phone", "[TELÉFONO]"),
            "direccion_completa": user_profile.get("address", "[DIRECCIÓN]"),
            "ciudad": user_profile.get("city", "[CIUDAD]")
        })
        
        # 2. Campos de entidades extraídas
        entities = extracted_entities.get("entities", [])
        for field_name, extractor in self.field_extractors.items():
            value = extractor(entities)
            if value:
                fields[field_name] = value
        
        # 3. Campos específicos del tipo de documento
        specific_data = extracted_entities.get("specific_data", {})
        fields.update(specific_data)
        
        # 4. Inferir tipo de cliente
        if fields.get("dni"):
            fields["tipo_cliente"] = "particular"
            fields["dni_cif"] = fields["dni"]
        elif fields.get("cif"):
            fields["tipo_cliente"] = "empresa"
            fields["dni_cif"] = fields["cif"]
        else:
            fields["tipo_cliente"] = "particular"
            fields["dni_cif"] = "[DNI/CIF]"
        
        # 5. Campos derivados del contexto de conversación
        fields.update(self._extract_context_fields(conversation_context))
        
        return fields
    
    def _extract_context_fields(self, conversation_context: str) -> Dict[str, Any]:
        """Extrae campos del contexto de la conversación usando IA."""
        try:
            prompt = f"""Analiza esta conversación y extrae información relevante para generar un documento legal:

CONVERSACIÓN:
{conversation_context[:2000]}

Extrae y devuelve SOLO un JSON con los siguientes campos si están presentes:
- tipo_reclamacion: "importe_excesivo", "servicios_no_prestados", "calidad_deficiente", etc.
- motivo_personalizado: descripción específica del problema
- empresa_destinataria: nombre de la empresa a la que se reclama
- organismo_regulador: organismo competente para el sector
- normativa_sectorial_especifica: normativa específica aplicable
- incluir_advertencia: true/false
- solicitud_adicional: petición adicional específica

FORMATO DE RESPUESTA:
{{"campo": "valor"}}

Responde SOLO con el JSON, sin explicaciones."""

            llm = get_llm(temperature=0.1, max_output_tokens=1000)
            response = llm.invoke(prompt)
            
            content = response.content if hasattr(response, 'content') else str(response)
            
            # Extraer JSON de la respuesta
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                try:
                    return json.loads(json_match.group())
                except json.JSONDecodeError:
                    pass
            
        except Exception as e:
            print(f"Error extracting context fields: {e}")
        
        return {}
    
    def _extract_entity_value(self, entities: List[Dict], entity_type: str) -> Optional[str]:
        """Extrae el valor de una entidad específica."""
        for entity in entities:
            if entity_type.lower() in entity.get("type", "").lower():
                if entity.get("confidence", 0) > 0.6:  # Solo entidades con alta confianza
                    return entity.get("value", "")
        return None
    
    def _evaluate_template_conditions(
        self,
        template: SmartTemplate,
        fields: Dict[str, Any],
        user_context: Dict[str, Any]
    ) -> List[TemplateSection]:
        """Evalúa condiciones y filtra secciones activas."""
        active_sections = []
        
        for section in template.sections:
            should_include = section.required  # Las secciones requeridas siempre se incluyen
            
            if not should_include and section.conditions:
                # Evaluar todas las condiciones (AND lógico)
                should_include = all(
                    self._evaluate_condition(condition, fields, user_context)
                    for condition in section.conditions
                )
            elif not should_include and not section.conditions:
                # Secciones sin condiciones se incluyen por defecto
                should_include = True
            
            if should_include:
                active_sections.append(section)
        
        # Ordenar por orden especificado
        active_sections.sort(key=lambda s: s.order)
        return active_sections
    
    def _evaluate_condition(
        self,
        condition: TemplateCondition,
        fields: Dict[str, Any],
        user_context: Dict[str, Any]
    ) -> bool:
        """Evalúa una condición individual."""
        field_value = fields.get(condition.field)
        
        if condition.operator == "exists":
            return field_value is not None and str(field_value).strip() != ""
        
        elif condition.operator == "not_exists":
            return field_value is None or str(field_value).strip() == ""
        
        elif condition.operator == "equals":
            return field_value == condition.value
        
        elif condition.operator == "contains":
            return condition.value in str(field_value) if field_value else False
        
        elif condition.operator == "greater_than":
            try:
                return float(field_value) > float(condition.value)
            except (ValueError, TypeError):
                return False
        
        elif condition.operator == "less_than":
            try:
                return float(field_value) < float(condition.value)
            except (ValueError, TypeError):
                return False
        
        return False
    
    def _enhance_sections_with_ai(
        self,
        sections: List[TemplateSection],
        fields: Dict[str, Any],
        conversation_context: str,
        user_context: Dict[str, Any]
    ) -> List[TemplateSection]:
        """Mejora las secciones con IA para mayor precisión."""
        enhanced_sections = []
        
        for section in sections:
            try:
                enhanced_content = self._enhance_section_content(
                    section, fields, conversation_context, user_context
                )
                
                enhanced_section = TemplateSection(
                    id=section.id,
                    title=section.title,
                    content=enhanced_content,
                    conditions=section.conditions,
                    required=section.required,
                    order=section.order
                )
                
                enhanced_sections.append(enhanced_section)
                
            except Exception as e:
                print(f"Error enhancing section {section.id}: {e}")
                enhanced_sections.append(section)  # Usar versión original si falla
        
        return enhanced_sections
    
    def _enhance_section_content(
        self,
        section: TemplateSection,
        fields: Dict[str, Any],
        conversation_context: str,
        user_context: Dict[str, Any]
    ) -> str:
        """Mejora el contenido de una sección específica con IA."""
        
        # Solo mejorar secciones que se beneficien de IA
        ai_enhanced_sections = [
            "motivo_importe_excesivo", "motivo_servicios_no_prestados",
            "alegacion_fondo", "alegaciones_especificas"
        ]
        
        if section.id not in ai_enhanced_sections:
            return section.content  # Retornar contenido original
        
        try:
            prompt = f"""Mejora este contenido de documento legal haciéndolo más específico y persuasivo:

SECCIÓN: {section.title}
CONTENIDO BASE:
{section.content}

CONTEXTO DE LA CONVERSACIÓN:
{conversation_context[:1500]}

CAMPOS DISPONIBLES:
{json.dumps({k: v for k, v in fields.items() if v and str(v) != '[CAMPO]'}, indent=2)}

INSTRUCCIONES:
1. Mantén la estructura y formato del contenido base
2. Añade detalles específicos basados en el contexto
3. Usa un lenguaje legal apropiado pero claro
4. Conserva todos los marcadores de campo {{campo}}
5. Haz el contenido más persuasivo y específico

Responde SOLO con el contenido mejorado, sin explicaciones."""

            llm = get_llm(temperature=0.3, max_output_tokens=1500)
            response = llm.invoke(prompt)
            
            enhanced_content = response.content if hasattr(response, 'content') else str(response)
            
            # Verificar que el contenido mejorado tenga sentido
            if len(enhanced_content.strip()) > 50 and '{' in enhanced_content:
                return enhanced_content.strip()
            
        except Exception as e:
            print(f"Error enhancing section content: {e}")
        
        return section.content
    
    def _generate_word_document(
        self,
        template: SmartTemplate,
        sections: List[TemplateSection],
        fields: Dict[str, Any],
        user_id: str
    ) -> str:
        """Genera documento Word con formato profesional."""
        
        # Crear documento
        doc = Document()
        
        # Configurar estilos
        self._setup_document_styles(doc)
        
        # Título del documento
        title = doc.add_heading(template.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generar secciones
        for section in sections:
            # Añadir título de sección si no es header o cierre
            if section.id not in ["header", "cierre"]:
                doc.add_heading(section.title, level=2)
            
            # Procesar contenido de la sección
            content = self._populate_template_fields(section.content, fields)
            
            # Añadir párrafos
            paragraphs = content.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    p = doc.add_paragraph(paragraph_text.strip())
                    
                    # Formato especial para ciertos tipos de contenido
                    if paragraph_text.startswith('ASUNTO:'):
                        p.style = 'Heading 3'
                    elif paragraph_text.startswith(('PRIMERA.-', 'SEGUNDA.-', 'TERCERA.-')):
                        p.style = 'Heading 4'
                    elif paragraph_text.strip().endswith(':'):
                        run = p.runs[0]
                        run.bold = True
            
            # Añadir espacio entre secciones
            doc.add_paragraph()
        
        # Guardar documento
        filename = f"{template.id}_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc_path = f"/tmp/{filename}"
        doc.save(doc_path)
        
        return doc_path
    
    def _setup_document_styles(self, doc: Document):
        """Configura estilos del documento."""
        try:
            # Estilo para encabezados
            heading_style = doc.styles['Heading 2']
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(14)
            heading_font.bold = True
            
            # Estilo para texto normal
            normal_style = doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Arial'
            normal_font.size = Pt(11)
            
        except Exception as e:
            print(f"Error setting up document styles: {e}")
    
    def _populate_template_fields(self, content: str, fields: Dict[str, Any]) -> str:
        """Rellena los campos de plantilla con valores reales."""
        result = content
        
        # Reemplazar campos con valores
        for field_name, field_value in fields.items():
            placeholder = f"{{{field_name}}}"
            if placeholder in result:
                # Convertir valor a string y limpiar
                str_value = str(field_value) if field_value is not None else f"[{field_name.upper()}]"
                result = result.replace(placeholder, str_value)
        
        # Limpiar campos no encontrados
        remaining_fields = re.findall(r'\{([^}]+)\}', result)
        for field in remaining_fields:
            placeholder = f"{{{field}}}"
            result = result.replace(placeholder, f"[{field.upper().replace('_', ' ')}]")
        
        return result
    
    def get_available_templates(self) -> List[Dict[str, Any]]:
        """Obtiene lista de plantillas disponibles."""
        return [
            {
                "id": template.id,
                "name": template.name,
                "description": template.description,
                "document_type": template.document_type,
                "sections_count": len(template.sections),
                "ai_enhanced": template.ai_enhancement
            }
            for template in self.templates.values()
        ]
    
    def get_template_preview(self, template_id: str, sample_fields: Dict[str, Any]) -> Dict[str, Any]:
        """Genera vista previa de una plantilla con campos de ejemplo."""
        template = self.templates.get(template_id)
        if not template:
            raise ValueError(f"Plantilla no encontrada: {template_id}")
        
        # Evaluar secciones con campos de ejemplo
        active_sections = self._evaluate_template_conditions(template, sample_fields, {})
        
        preview_sections = []
        for section in active_sections:
            content_preview = self._populate_template_fields(section.content, sample_fields)
            preview_sections.append({
                "title": section.title,
                "content_preview": content_preview[:300] + "..." if len(content_preview) > 300 else content_preview,
                "required": section.required,
                "order": section.order
            })
        
        return {
            "template_name": template.name,
            "sections": preview_sections,
            "total_sections": len(active_sections),
            "ai_enhancement": template.ai_enhancement
        }


# Instancia global del servicio
intelligent_template_service = IntelligentTemplateService()