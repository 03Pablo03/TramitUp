import io
from datetime import datetime
from docx import Document
from docx.shared import Pt

from app.ai.rag.retriever import retrieve_context
from app.ai.chains.classify_chain import classify_tramite
from app.ai.chains.document_chain import generate_document
from app.ai.prompts.system_base import AVISO_LEGAL_UI


def _html_template(title: str, content: str) -> str:
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: 'Georgia', serif; margin: 2cm; line-height: 1.6; }}
h1 {{ font-size: 18pt; margin-bottom: 1em; }}
.aviso {{ font-size: 9pt; color: #666; margin-top: 2em; padding: 1em; border: 1px solid #ccc; }}
</style>
</head>
<body>
<h1>{title}</h1>
<div>{content.replace(chr(10), '<br>')}</div>
<p class="aviso">{AVISO_LEGAL_UI}</p>
<p style="margin-top: 2em;">Fecha: {datetime.now().strftime('%d/%m/%Y')}</p>
</body>
</html>"""


def generate_pdf(title: str, content: str) -> bytes:
    """Generate PDF from HTML content. Requires WeasyPrint + GTK3 on Windows."""
    try:
        from weasyprint import HTML
    except OSError as e:
        if "gobject" in str(e).lower() or "pango" in str(e).lower():
            raise RuntimeError(
                "WeasyPrint necesita GTK3 en Windows. Instala: "
                "https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases"
            ) from e
        raise
    html_str = _html_template(title, content)
    html = HTML(string=html_str)
    buffer = io.BytesIO()
    html.write_pdf(buffer)
    return buffer.getvalue()


def generate_docx(title: str, content: str) -> bytes:
    """Generate Word document."""
    doc = Document()
    doc.add_heading(title, 0)
    for para in content.split("\n\n"):
        if para.strip():
            p = doc.add_paragraph(para.replace("\n", " "))
            p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()
    doc.add_paragraph(AVISO_LEGAL_UI).alignment = 1  # center-like
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_document_from_chat(
    user_id: str,
    conversation_id: str | None,
    case_description: str,
    doc_type: str,
) -> tuple[str, str]:
    """
    Create document: RAG + classify + generate. Returns (document_id, content).
    """
    from app.core.supabase_client import get_supabase_client

    classification = classify_tramite(case_description)
    rag_chunks = retrieve_context(case_description)
    context = "\n\n".join(c["content"] for c in rag_chunks)
    content = generate_document(context, case_description, doc_type)
    title = classification.get("titulo_resumen", "Documento")[:200]
    category = classification.get("category", classification.get("categoria", "otro"))

    supabase = get_supabase_client()
    result = supabase.table("documents").insert({
        "user_id": user_id,
        "conversation_id": conversation_id,
        "title": title,
        "content": content,
        "category": category,
    }).execute()
    doc_id = result.data[0]["id"]
    return doc_id, content
