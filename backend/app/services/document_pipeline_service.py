"""
Pipeline de generación de documentos: extracción → generación → renderizado → storage.
"""
import io
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.supabase_client import get_supabase_client
from app.ai.chains.document_extract_chain import extract_document_data
from app.ai.chains.document_generate_chain import generate_document_content



def _render_pdf(doc: dict) -> bytes:
    """Genera PDF usando ReportLab (puro Python, sin dependencias nativas)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT

    enc = doc.get("encabezado", {})
    cuerpo = doc.get("cuerpo", {})
    titulo = doc.get("titulo", "Modelo de escrito")
    disclaimer = doc.get("nota_disclaimer", "")

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=3 * cm, rightMargin=3 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    normal = ParagraphStyle("normal", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=8)
    bold_style = ParagraphStyle("bold", parent=normal, fontName="Helvetica-Bold")
    heading = ParagraphStyle("heading", parent=styles["Heading1"], fontSize=14, spaceAfter=16, alignment=TA_CENTER)
    justified = ParagraphStyle("justified", parent=normal, alignment=TA_JUSTIFY)
    small = ParagraphStyle("small", parent=normal, fontSize=8, textColor=colors.grey, alignment=TA_CENTER)

    story = []

    # Título
    story.append(Paragraph(titulo, heading))
    story.append(Spacer(1, 0.3 * cm))

    # Remitente / destinatario / lugar
    for block in [enc.get("remitente_bloque"), enc.get("destinatario_bloque"), enc.get("lugar_fecha")]:
        if block:
            for line in block.split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), normal))
            story.append(Spacer(1, 0.3 * cm))

    # Asunto
    if enc.get("asunto"):
        story.append(Paragraph(f"<b>ASUNTO: {enc['asunto']}</b>", bold_style))
        story.append(Spacer(1, 0.4 * cm))

    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
    story.append(Spacer(1, 0.4 * cm))

    # Cuerpo
    if cuerpo.get("saludo"):
        story.append(Paragraph(cuerpo["saludo"], normal))
        story.append(Spacer(1, 0.2 * cm))

    for p in cuerpo.get("parrafos_hechos", []):
        if p:
            story.append(Paragraph(p, justified))

    for p in cuerpo.get("parrafos_fundamentos", []):
        if p:
            story.append(Paragraph(p, justified))

    if cuerpo.get("parrafo_solicitud"):
        story.append(Paragraph(cuerpo["parrafo_solicitud"], justified))

    if cuerpo.get("cierre"):
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph(cuerpo["cierre"], normal))

    # Firma
    if doc.get("firma_bloque"):
        story.append(Spacer(1, 1 * cm))
        for line in doc["firma_bloque"].split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), normal))

    # Disclaimer
    if disclaimer:
        story.append(Spacer(1, 0.8 * cm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(disclaimer, small))

    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Modelo generado por Tramitup · tramitup.com", small))

    pdf.build(story)
    return buffer.getvalue()


def _render_docx(doc: dict) -> bytes:
    document = Document()
    enc = doc.get("encabezado", {})
    cuerpo = doc.get("cuerpo", {})
    document.add_heading(doc.get("titulo", "Modelo de escrito"), 0)
    for block in [enc.get("remitente_bloque"), enc.get("destinatario_bloque"), enc.get("lugar_fecha")]:
        if block:
            p = document.add_paragraph(block.replace("\n", " "))
            p.paragraph_format.space_after = Pt(6)
    if enc.get("asunto"):
        p = document.add_paragraph(enc["asunto"])
        p.bold = True
        p.paragraph_format.space_after = Pt(12)
    if cuerpo.get("saludo"):
        document.add_paragraph(cuerpo["saludo"])
    for p in cuerpo.get("parrafos_hechos", []):
        document.add_paragraph(p)
    for p in cuerpo.get("parrafos_fundamentos", []):
        document.add_paragraph(p)
    if cuerpo.get("parrafo_solicitud"):
        document.add_paragraph(cuerpo["parrafo_solicitud"])
    if cuerpo.get("cierre"):
        document.add_paragraph(cuerpo["cierre"])
    if doc.get("firma_bloque"):
        document.add_paragraph(doc["firma_bloque"])
    if doc.get("nota_disclaimer"):
        dp = document.add_paragraph(doc["nota_disclaimer"])
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def preview_document_content(user_id: str, conversation_id: str) -> dict:
    """
    Genera preview del documento sin renderizar a PDF/DOCX.
    Retorna el contenido estructurado para visualizar en el frontend.
    """
    supabase = get_supabase_client()
    result = supabase.table("messages").select("role, content").eq(
        "conversation_id", conversation_id
    ).order("created_at", desc=False).execute()
    messages = [{"role": r["role"], "content": r["content"]} for r in (result.data or [])]
    if not messages:
        raise ValueError("Conversación vacía o no encontrada")
    conv = supabase.table("conversations").select("user_id").eq(
        "id", conversation_id
    ).eq("user_id", user_id).execute()
    if not conv.data:
        raise ValueError("Conversación no encontrada o no autorizada")

    extracted = extract_document_data(messages)
    doc_content = generate_document_content(extracted)

    # Build a plain text preview from the structured content
    enc = doc_content.get("encabezado", {})
    cuerpo = doc_content.get("cuerpo", {})
    sections = []

    if enc.get("remitente_bloque"):
        sections.append({"type": "header", "label": "Remitente", "content": enc["remitente_bloque"]})
    if enc.get("destinatario_bloque"):
        sections.append({"type": "header", "label": "Destinatario", "content": enc["destinatario_bloque"]})
    if enc.get("lugar_fecha"):
        sections.append({"type": "header", "label": "Lugar y fecha", "content": enc["lugar_fecha"]})
    if enc.get("asunto"):
        sections.append({"type": "subject", "label": "Asunto", "content": enc["asunto"]})
    if cuerpo.get("saludo"):
        sections.append({"type": "body", "label": "Saludo", "content": cuerpo["saludo"]})
    for i, p in enumerate(cuerpo.get("parrafos_hechos", [])):
        if p:
            sections.append({"type": "body", "label": f"Hechos ({i + 1})", "content": p})
    for i, p in enumerate(cuerpo.get("parrafos_fundamentos", [])):
        if p:
            sections.append({"type": "body", "label": f"Fundamentos ({i + 1})", "content": p})
    if cuerpo.get("parrafo_solicitud"):
        sections.append({"type": "body", "label": "Solicitud", "content": cuerpo["parrafo_solicitud"]})
    if cuerpo.get("cierre"):
        sections.append({"type": "body", "label": "Cierre", "content": cuerpo["cierre"]})
    if doc_content.get("firma_bloque"):
        sections.append({"type": "signature", "label": "Firma", "content": doc_content["firma_bloque"]})

    return {
        "title": doc_content.get("titulo", "Modelo de escrito"),
        "document_type": extracted.get("document_type", "solicitud"),
        "sections": sections,
        "disclaimer": doc_content.get("nota_disclaimer", ""),
    }


def run_document_pipeline(
    user_id: str,
    conversation_id: str,
    format_request: str,  # "pdf" | "docx" | "both"
) -> dict:
    """
    Ejecuta el pipeline completo.
    Returns: { document_id, pdf_url?, docx_url?, document_type, generated_at }
    """
    supabase = get_supabase_client()
    result = supabase.table("messages").select("role, content").eq(
        "conversation_id", conversation_id
    ).order("created_at", desc=False).execute()
    messages = [{"role": r["role"], "content": r["content"]} for r in (result.data or [])]
    if not messages:
        raise ValueError("Conversación vacía o no encontrada")
    conv = supabase.table("conversations").select("user_id").eq("id", conversation_id).eq("user_id", user_id).execute()
    if not conv.data:
        raise ValueError("Conversación no encontrada o no autorizada")
    extracted = extract_document_data(messages)
    doc_content = generate_document_content(extracted)
    doc_type = extracted.get("document_type", "solicitud_informacion_administrativa")
    doc_id = str(uuid.uuid4())
    pdf_path = None
    docx_path = None
    bucket = "documents"
    pdf_bytes = _render_pdf(doc_content) if format_request in ("pdf", "both") else None
    docx_bytes = _render_docx(doc_content) if format_request in ("docx", "both") else None
    try:
        if pdf_bytes:
            pdf_path = f"{user_id}/{doc_id}/documento.pdf"
            supabase.storage.from_(bucket).upload(
                file=pdf_bytes,
                path=pdf_path,
                file_options={"content-type": "application/pdf", "upsert": "true"},
            )
        if docx_bytes:
            docx_path = f"{user_id}/{doc_id}/documento.docx"
            supabase.storage.from_(bucket).upload(
                file=docx_bytes,
                path=docx_path,
                file_options={
                    "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "upsert": "true",
                },
            )
        supabase.table("generated_documents").insert({
            "id": doc_id,
            "user_id": user_id,
            "conversation_id": conversation_id,
            "document_type": doc_type,
            "data_json": extracted,
            "pdf_path": pdf_path,
            "docx_path": docx_path,
        }).execute()
    except Exception as e:
        if "does not exist" in str(e).lower() or "bucket" in str(e).lower():
            raise ValueError(
                "El bucket 'documents' no existe en Supabase Storage. "
                "Créalo en Dashboard > Storage > New bucket (privado)."
            ) from e
        raise
    expiry = 3600
    pdf_url = None
    docx_url = None

    def _signed_url(path: str) -> str | None:
        r = supabase.storage.from_(bucket).create_signed_url(path, expiry)
        return r.get("signedUrl") or r.get("signedURL") or (r.get("data") or {}).get("signedUrl")

    if pdf_path:
        pdf_url = _signed_url(pdf_path)
    if docx_path:
        docx_url = _signed_url(docx_path)
    return {
        "document_id": doc_id,
        "pdf_url": pdf_url,
        "docx_url": docx_url,
        "document_type": doc_type,
        "generated_at": datetime.utcnow().isoformat() + "Z",
    }


def get_document_download_urls(document_id: str, user_id: str) -> dict:
    """Genera URLs firmadas frescas para un documento."""
    supabase = get_supabase_client()
    r = supabase.table("generated_documents").select("pdf_path, docx_path").eq("id", document_id).eq("user_id", user_id).execute()
    if not r.data:
        raise ValueError("Documento no encontrado")
    row = r.data[0]
    pdf_path = row.get("pdf_path")
    docx_path = row.get("docx_path")
    expiry = 3600
    def _signed(p: str | None) -> str | None:
        if not p:
            return None
        resp = supabase.storage.from_("documents").create_signed_url(p, expiry)
        return resp.get("signedUrl") or resp.get("signedURL") or (resp.get("data") or {}).get("signedUrl")

    pdf_url = _signed(pdf_path)
    docx_url = _signed(docx_path)
    return {"pdf_url": pdf_url, "docx_url": docx_url}


def get_document_history(user_id: str) -> list[dict]:
    """Lista documentos generados por el usuario."""
    supabase = get_supabase_client()
    r = supabase.table("generated_documents").select(
        "id, conversation_id, document_type, pdf_path, docx_path, created_at"
    ).eq("user_id", user_id).order("created_at", desc=True).limit(100).execute()
    def _to_item(d: dict) -> dict:
        ca = d.get("created_at")
        return {
            "document_id": d["id"],
            "conversation_id": d.get("conversation_id"),
            "document_type": d["document_type"],
            "created_at": ca.isoformat() if hasattr(ca, "isoformat") else str(ca) if ca else "",
            "has_pdf": bool(d.get("pdf_path")),
            "has_docx": bool(d.get("docx_path")),
        }

    return [_to_item(d) for d in (r.data or [])]
